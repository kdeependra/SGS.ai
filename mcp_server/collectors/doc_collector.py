"""Document metadata collector – extracts metadata from text documents."""
import io
import os
import re
from collections import Counter


_SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".xml", ".html", ".yaml", ".yml", ".log", ".rst", ".csv", ".docx", ".pdf"}


def _detect_doc_type(ext: str, content: str) -> str:
    mapping = {
        ".md": "markdown", ".json": "json", ".xml": "xml",
        ".html": "html", ".yaml": "yaml", ".yml": "yaml",
        ".log": "log", ".rst": "restructuredtext", ".csv": "csv",
        ".txt": "plaintext", ".docx": "docx", ".pdf": "pdf",
    }
    return mapping.get(ext, "plaintext")


def _infer_type(values: list[str]) -> str:
    """Infer data type from a list of sample values."""
    int_count = float_count = bool_count = 0
    for v in values:
        v = v.strip()
        if not v or v.lower() in ("null", "none", ""):
            continue
        if v.lower() in ("true", "false"):
            bool_count += 1
            continue
        try:
            int(v)
            int_count += 1
            continue
        except ValueError:
            pass
        try:
            float(v)
            float_count += 1
            continue
        except ValueError:
            pass
    total = len([x for x in values if x.strip() and x.strip().lower() not in ("null", "none")])
    if total == 0:
        return "string"
    if int_count == total:
        return "integer"
    if (int_count + float_count) == total:
        return "float"
    if bool_count == total:
        return "boolean"
    return "string"


def _extract_markdown_tables(content: str) -> list[dict]:
    """Extract tables from markdown pipe-delimited format."""
    tables = []
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if "|" in line and line.startswith("|"):
            headers = [c.strip() for c in line.strip("|").split("|")]
            headers = [h for h in headers if h]
            if i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
                rows = []
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                    rows.append(cells)
                    j += 1
                columns = []
                for ci, h in enumerate(headers):
                    col_values = [r[ci] if ci < len(r) else "" for r in rows]
                    columns.append({
                        "name": h,
                        "data_type": _infer_type(col_values),
                        "sample_values": col_values[:5],
                        "nullable": any(v.strip() == "" for v in col_values),
                    })
                tables.append({
                    "table": f"table_{len(tables) + 1}",
                    "row_count": len(rows),
                    "columns": columns,
                })
                i = j
                continue
        i += 1
    return tables


def _extract_html_tables(content: str) -> list[dict]:
    """Extract tables from HTML <table> elements."""
    tables = []
    table_blocks = re.findall(r"<table[^>]*>(.*?)</table>", content, re.DOTALL | re.IGNORECASE)
    for block in table_blocks:
        headers = re.findall(r"<th[^>]*>(.*?)</th>", block, re.DOTALL | re.IGNORECASE)
        headers = [re.sub(r"<[^>]+>", "", h).strip() for h in headers]
        row_matches = re.findall(r"<tr[^>]*>(.*?)</tr>", block, re.DOTALL | re.IGNORECASE)
        rows = []
        for rm in row_matches:
            cells = re.findall(r"<td[^>]*>(.*?)</td>", rm, re.DOTALL | re.IGNORECASE)
            cells = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
            if cells:
                rows.append(cells)
        if not headers and rows:
            headers = rows[0]
            rows = rows[1:]
        if headers:
            columns = []
            for ci, h in enumerate(headers):
                col_values = [r[ci] if ci < len(r) else "" for r in rows]
                columns.append({
                    "name": h,
                    "data_type": _infer_type(col_values),
                    "sample_values": col_values[:5],
                    "nullable": any(v.strip() == "" for v in col_values),
                })
            tables.append({
                "table": f"table_{len(tables) + 1}",
                "row_count": len(rows),
                "columns": columns,
            })
    return tables


def _extract_plaintext_tables(content: str) -> list[dict]:
    """Detect delimiter-separated tabular data in plaintext/log files.

    Supports multiple tables separated by blank or non-tabular lines.
    Handles tab, pipe, semicolon, comma and multi-space delimiters.
    """
    import re
    all_lines = content.splitlines()
    if len(all_lines) < 2:
        return []

    def _split_multispace(line: str) -> list[str]:
        """Split on 2+ consecutive spaces (for space-aligned PDF text)."""
        return [c.strip() for c in re.split(r'  +', line.strip()) if c.strip()]

    # Try common delimiters; multi-space last (most permissive)
    delimiters: list[tuple[str | None, str]] = [
        ("\t", "tab"), ("|", "pipe"), (";", "semicolon"), (",", "comma"),
        (None, "multispace"),  # None signals multi-space split
    ]
    for delim, delim_name in delimiters:
        tables: list[dict] = []
        consumed: set[int] = set()  # line indices already part of a table
        for start in range(len(all_lines)):
            if start in consumed:
                continue
            line = all_lines[start].strip()
            if not line:
                continue
            if delim is not None:
                header_parts = [c.strip() for c in line.split(delim)]
            else:
                header_parts = _split_multispace(line)
            if len(header_parts) < 2:
                continue
            expected = len(header_parts)
            # Collect contiguous data rows right after the header
            rows: list[list[str]] = []
            end = start + 1
            while end < len(all_lines):
                row_line = all_lines[end].strip()
                if not row_line:
                    # blank line ends the current table block
                    break
                if delim is not None:
                    cells = [c.strip() for c in row_line.split(delim)]
                else:
                    cells = _split_multispace(row_line)
                if len(cells) != expected:
                    break
                rows.append(cells)
                end += 1
            if len(rows) < 3:
                continue
            # Build columns with inferred types
            columns = []
            for ci, h in enumerate(header_parts):
                col_values = [r[ci] if ci < len(r) else "" for r in rows]
                columns.append({
                    "name": h,
                    "data_type": _infer_type(col_values),
                    "sample_values": col_values[:5],
                    "nullable": any(v.strip() == "" for v in col_values),
                })
            tables.append({"table": f"table_{len(tables)+1}", "row_count": len(rows), "columns": columns})
            # Mark header + data lines as consumed so we don't re-detect them
            for idx in range(start, end):
                consumed.add(idx)
        if tables:
            return tables
    return []


def _extract_docx_text_and_tables(raw: bytes) -> tuple[str, list[dict], dict]:
    """Extract full text, tables, and document properties from a .docx file."""
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(raw))
    # Full text from paragraphs + table cells
    text_parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            text_parts.append(" ".join(cell.text.strip() for cell in row.cells))
    text = "\n".join(text_parts)
    tables: list[dict] = []
    for tbl in doc.tables:
        rows_data: list[list[str]] = []
        for row in tbl.rows:
            rows_data.append([cell.text.strip() for cell in row.cells])
        if not rows_data:
            continue
        headers = rows_data[0]
        data_rows = rows_data[1:]
        columns = []
        for ci, h in enumerate(headers):
            col_values = [r[ci] if ci < len(r) else "" for r in data_rows]
            columns.append({
                "name": h or f"col_{ci+1}",
                "data_type": _infer_type(col_values),
                "sample_values": col_values[:5],
                "nullable": any(v.strip() == "" for v in col_values),
            })
        tables.append({"table": f"table_{len(tables)+1}", "row_count": len(data_rows), "columns": columns})

    # Extract document properties
    props: dict = {}
    try:
        cp = doc.core_properties
        for attr in ("author", "title", "subject", "keywords", "category",
                     "comments", "last_modified_by", "revision", "version",
                     "content_status", "language", "identifier"):
            val = getattr(cp, attr, None)
            if val:
                props[attr] = str(val)
        if cp.created:
            props["created"] = str(cp.created)
        if cp.modified:
            props["modified"] = str(cp.modified)
    except Exception:
        pass
    props["paragraph_count"] = len(doc.paragraphs)
    props["table_count"] = len(doc.tables)
    sections = list(doc.sections)
    props["section_count"] = len(sections)

    return text, tables, props


def _extract_pdf_text_and_tables(raw: bytes) -> tuple[str, list[dict], dict]:
    """Extract full text, tables, and document properties from a PDF file."""
    import pdfplumber
    import re as _re
    text_parts: list[str] = []
    tables: list[dict] = []
    props: dict = {}
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        # Extract PDF metadata
        meta = pdf.metadata or {}
        for key in ("Author", "Title", "Subject", "Keywords", "Creator",
                    "Producer", "CreationDate", "ModDate"):
            val = meta.get(key)
            if val:
                props[key.lower()] = str(val)
        props["page_count"] = len(pdf.pages)

        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                page_text = _re.sub(r'\(cid:415\)', 'ti', page_text)
                page_text = _re.sub(r'\(cid:\d+\)', '', page_text)
                text_parts.append(page_text)
            raw_tables = page.extract_tables() or []
            for tbl in raw_tables:
                if tbl:
                    for row in tbl:
                        text_parts.append(" ".join((c or "").strip() for c in row))
            for tbl in raw_tables:
                if not tbl or len(tbl) < 2:
                    continue
                headers = [c or f"col_{i+1}" for i, c in enumerate(tbl[0])]
                data_rows = tbl[1:]
                columns = []
                for ci, h in enumerate(headers):
                    col_values = [(r[ci] or "").strip() if ci < len(r) else "" for r in data_rows]
                    columns.append({
                        "name": h.strip(),
                        "data_type": _infer_type(col_values),
                        "sample_values": col_values[:5],
                        "nullable": any(v == "" for v in col_values),
                    })
                tables.append({"table": f"table_{len(tables)+1}", "row_count": len(data_rows), "columns": columns})

        # If no native tables found, try word-position extraction across all pages
        if not tables:
            positional_tables = _extract_pdf_tables_by_position(pdf)
            if positional_tables:
                tables = positional_tables

    props["table_count"] = len(tables)
    return "\n".join(text_parts), tables, props


def _extract_pdf_tables_by_position(pdf) -> list[dict]:
    """Extract tables from PDF by analysing word bounding-box positions.

    Works for 'Print to PDF' documents where text wraps across lines but
    the original column positions are preserved spatially.  Handles records
    that span multiple visual lines by trying adaptive merge factors.
    """
    import re as _re
    all_tables: list[dict] = []

    for page in pdf.pages:
        words = page.extract_words(x_tolerance=3, y_tolerance=3)
        if not words:
            continue

        # Clean CID artifacts in word text
        for w in words:
            w["text"] = _re.sub(r'\(cid:415\)', 'ti', w["text"])
            w["text"] = _re.sub(r'\(cid:\d+\)', '', w["text"])

        # Group words into visual rows by y-coordinate (tolerance 4px)
        words.sort(key=lambda w: (w["top"], w["x0"]))
        y_lines: list[list[dict]] = []
        cur_group = [words[0]]
        for w in words[1:]:
            if abs(w["top"] - cur_group[0]["top"]) < 4:
                cur_group.append(w)
            else:
                y_lines.append(sorted(cur_group, key=lambda x: x["x0"]))
                cur_group = [w]
        y_lines.append(sorted(cur_group, key=lambda x: x["x0"]))

        # Build tab-separated text for each visual line.
        tab_lines: list[str] = []
        for row_words in y_lines:
            parts: list[str] = []
            prev_x1 = None
            for w in row_words:
                if prev_x1 is not None:
                    gap = w["x0"] - prev_x1
                    if gap > 12:
                        parts.append("\t")
                    elif gap > 0:
                        parts.append(" ")
                parts.append(w["text"])
                prev_x1 = w["x1"]
            line = "".join(parts).strip()
            if line:
                tab_lines.append(line)

        # Adaptive table detection
        tables = _detect_tables_adaptive(tab_lines)
        if tables:
            all_tables.extend(tables)

    return all_tables


def _detect_tables_adaptive(tab_lines: list[str]) -> list[dict]:
    """Detect tables in tab-separated lines, trying different merge factors.

    For each candidate header position, try reading 1, 2, or 3 visual lines
    as one record.  Whichever factor yields ≥2 consistent data rows wins.
    """
    tables: list[dict] = []
    consumed: set[int] = set()

    for start in range(len(tab_lines)):
        if start in consumed:
            continue

        best_table = None
        best_end = start

        for mf in (1, 2, 3):
            if start + mf > len(tab_lines):
                continue
            header = "\t".join(tab_lines[start:start + mf])
            header_parts = [c.strip() for c in header.split("\t")]
            if len(header_parts) < 3:
                continue
            expected = len(header_parts)

            rows: list[list[str]] = []
            end = start + mf
            while end + mf <= len(tab_lines):
                record = "\t".join(tab_lines[end:end + mf])
                cells = [c.strip() for c in record.split("\t")]
                if len(cells) != expected:
                    break
                rows.append(cells)
                end += mf

            if len(rows) >= 3:
                columns = []
                for ci, h in enumerate(header_parts):
                    col_values = [r[ci] if ci < len(r) else "" for r in rows]
                    columns.append({
                        "name": h,
                        "data_type": _infer_type(col_values),
                        "sample_values": col_values[:5],
                        "nullable": any(v.strip() == "" for v in col_values),
                    })
                candidate = {
                    "table": f"table_{len(tables) + 1}",
                    "row_count": len(rows),
                    "columns": columns,
                }
                if best_table is None or len(rows) > best_table["row_count"]:
                    best_table = candidate
                    best_end = end

        if best_table:
            tables.append(best_table)
            for idx in range(start, best_end):
                consumed.add(idx)

    return tables


def _extract_tables(content: str, doc_type: str) -> list[dict]:
    """Extract tabular data from document content based on doc_type."""
    if doc_type == "markdown":
        return _extract_markdown_tables(content)
    elif doc_type in ("html", "xml"):
        return _extract_html_tables(content)
    elif doc_type in ("plaintext", "log"):
        return _extract_plaintext_tables(content)
    return []


def _extract_structure(content: str, doc_type: str) -> dict:
    structure: dict = {}
    if doc_type == "markdown":
        headings = re.findall(r"^(#{1,6})\s+(.+)", content, re.MULTILINE)
        structure["headings"] = [{"level": len(h[0]), "text": h[1].strip()} for h in headings]
        structure["link_count"] = len(re.findall(r"\[.*?\]\(.*?\)", content))
        structure["code_block_count"] = content.count("```") // 2
    elif doc_type == "json":
        import json as _json
        try:
            parsed = _json.loads(content)
            if isinstance(parsed, dict):
                structure["top_level_keys"] = list(parsed.keys())[:50]
            elif isinstance(parsed, list):
                structure["array_length"] = len(parsed)
        except _json.JSONDecodeError:
            structure["parse_error"] = True
    elif doc_type in ("xml", "html"):
        tags = re.findall(r"<(\w+)[\s>]", content)
        structure["tag_counts"] = dict(Counter(tags).most_common(20))
    return structure


def collect_document_metadata(file_path: str | None = None, content: str | None = None,
                              file_name: str | None = None, raw_bytes: bytes | None = None) -> dict:
    """Analyse a document and return its metadata."""
    binary_tables: list[dict] = []
    doc_properties: dict = {}

    if file_path:
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        display_name = file_name or os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        if ext in (".docx", ".pdf"):
            with open(file_path, "rb") as f:
                raw_bytes = f.read()
            if ext == ".docx":
                text, binary_tables, doc_properties = _extract_docx_text_and_tables(raw_bytes)
            else:
                text, binary_tables, doc_properties = _extract_pdf_text_and_tables(raw_bytes)
        else:
            with open(file_path, encoding="utf-8", errors="replace") as f:
                text = f.read()
    elif raw_bytes is not None:
        display_name = file_name or "uploaded_document"
        file_size = len(raw_bytes)
        ext = os.path.splitext(display_name)[1].lower() if display_name else ""
        if ext == ".docx":
            text, binary_tables, doc_properties = _extract_docx_text_and_tables(raw_bytes)
        elif ext == ".pdf":
            text, binary_tables, doc_properties = _extract_pdf_text_and_tables(raw_bytes)
        else:
            text = raw_bytes.decode("utf-8", errors="replace")
    elif content is not None:
        text = content
        display_name = file_name or "uploaded_document"
        file_size = len(content.encode("utf-8"))
        ext = os.path.splitext(display_name)[1].lower() if display_name else ""
    else:
        raise ValueError("Provide either file_path, content, or raw_bytes")

    doc_type = _detect_doc_type(ext, text)
    lines = text.splitlines()
    words = text.split()

    word_freq = Counter(w.lower().strip(".,;:!?\"'()[]{}") for w in words if len(w) > 2)

    structure = _extract_structure(text, doc_type)
    # Merge document properties into structure
    if doc_properties:
        structure["document_properties"] = doc_properties

    # Extract tabular data from document content
    # For binary formats (docx/pdf): prefer native binary extraction (pdfplumber /
    # python-docx) since it uses actual document structure.  Supplement with
    # plaintext detection only for *additional* tables not already found.
    if doc_type in ("docx", "pdf"):
        if binary_tables:
            # Start with structurally-detected tables
            extracted_tables = list(binary_tables)
            # Try plaintext detection for any tables the binary extractor missed
            plaintext_tables = _extract_plaintext_tables(text)
            # Only add plaintext tables whose column signatures are genuinely new
            existing_sigs = {
                tuple(sorted(c["name"] for c in t["columns"]))
                for t in extracted_tables
            }
            for pt in plaintext_tables:
                sig = tuple(sorted(c["name"] for c in pt["columns"]))
                if sig not in existing_sigs:
                    extracted_tables.append(pt)
                    existing_sigs.add(sig)
        else:
            # Binary extractor found nothing — fall back to plaintext
            extracted_tables = _extract_plaintext_tables(text)
    else:
        extracted_tables = _extract_tables(text, doc_type)

    # Rename tables to filename-based names: filename_1, filename_2, ...
    if extracted_tables:
        base_name = display_name.rsplit(".", 1)[0] if "." in display_name else display_name
        for i, tbl in enumerate(extracted_tables, 1):
            tbl["table"] = f"{base_name}_{i}"

    # Update table_count in document_properties to reflect final extraction result
    if doc_properties and "table_count" in doc_properties:
        doc_properties["table_count"] = len(extracted_tables)
        structure["document_properties"] = doc_properties

    result = {
        "source_type": "document",
        "file_name": display_name,
        "file_size": file_size,
        "doc_type": doc_type,
        "line_count": len(lines),
        "word_count": len(words),
        "char_count": len(text),
        "avg_line_length": round(sum(len(l) for l in lines) / max(len(lines), 1), 1),
        "top_words": dict(word_freq.most_common(20)),
        "structure": structure,
    }

    if extracted_tables:
        result["tables"] = extracted_tables
        total_cols = sum(len(t.get("columns", [])) for t in extracted_tables)
        total_rows = sum(t.get("row_count", 0) for t in extracted_tables)
        result["column_count"] = total_cols
        result["row_count"] = total_rows

    return result
